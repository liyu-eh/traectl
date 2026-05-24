#!/usr/bin/env python3
"""生成 traectl 架构文档（.docx）—— 原版风格 + 当前源码数据"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 全局样式 ──────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

for level in range(1, 5):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Microsoft YaHei"
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

# ── 辅助函数 ──────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """给单元格设置背景色"""
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def add_table(headers, rows, col_widths=None, header_bg="1A5CB0"):
    """添加格式化表格——表头白字蓝底"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_bg)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            # 奇偶行交替背景
            if ri % 2 == 1:
                set_cell_shading(cell, "F2F6FC")
    if col_widths:
        for row_obj in table.rows:
            for ci, w in enumerate(col_widths):
                row_obj.cells[ci].width = Cm(w)
    doc.add_paragraph()  # 表后间距
    return table

def add_step_table(steps):
    return add_table(["#", "步骤说明"], [[str(i), s] for i, s in steps], col_widths=[1.2, 14.8])

def add_code_block(text):
    """添加代码块风格段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.name = "Consolas"
    return p

# ══════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("traectl-CLI 技术栈与架构文档")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x5C, 0xB0)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("版本 3.1.0 · Python ≥ 3.11 · 6116 行源码 · 172 测试 · CDP JavaScript 注入控制")
run.font.size = Pt(10.5)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # 间距

# ══════════════════════════════════════════════════════════════════
# 1. 📋 技术栈
# ══════════════════════════════════════════════════════════════════
doc.add_heading("1. 📋 技术栈", level=1)
doc.add_paragraph("以下表格列出 traectl-CLI 实际采用的核心技术及其用途。")

add_table(
    ["技术", "版本", "用途"],
    [
        ["Python", "≥ 3.11", "项目运行语言，充分利用 asyncio 异步 IO"],
        ["Typer (Click)", "≥ 0.12", "CLI 框架：参数解析、子命令注册、自动帮助文档"],
        ["websockets", "≥ 12", "CDP WebSocket 全双工通信，与 Chrome DevTools 协议交互"],
        ["Rich", "≥ 13", "人类可读输出：stderr 日志、Syntax 代码高亮、面板渲染"],
        ["pytest", "≥ 8", "单元测试框架，覆盖 CLI、CDP、配置、工作区等模块"],
        ["asyncio", "stdlib", "异步 IO 驱动：CDP 通信、连接池、响应轮询"],
        ["setuptools", "≥ 64", "项目打包与 editable install（pyproject.toml 声明）"],
    ],
    col_widths=[3, 1.8, 11.2],
)

p = doc.add_paragraph()
run = p.add_run("注：pyproject.toml 中显式依赖只有 typer 和 rich；websockets 为运行时直接导入，pytest 为测试依赖。")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ══════════════════════════════════════════════════════════════════
# 2. 🏗️ 三层架构
# ══════════════════════════════════════════════════════════════════
doc.add_heading("2. 🏗️ 三层架构", level=1)
doc.add_paragraph("traectl-CLI 采用清晰的三层架构设计，从顶层 CLI 到底层 CDP 通信逐层隔离职责：")

add_table(
    ["层", "模块", "职责"],
    [
        ["🖥️ CLI 层", "cli/ 7 子模块 (1177 行)", "Typer 命令定义、参数解析、输出格式化（JSON / --human / --output ndjson）"],
        ["🖥️ CLI 层", "response.py (176 行)", "标准化响应外壳：StandardResponse + JsonResponse + ok/error/dry_run 工厂"],
        ["🖥️ CLI 层", "config.py (235 行)", "常量/退出码/DOM 选择器（~30 个）/Agent 角色模板/环境变量 TRAECTL_* 前缀"],
        ["🎮 Controller 层", "controller.py (36 行)", "TraeSoloController——组合 8 个 Mixin 的主控制器"],
        ["🎮 Controller 层", "mixins/ 共 8 文件 (1,178 行)", "Chat / Task / Model / Editor / Terminal / Git / Media / Health"],
        ["🎮 Controller 层", "mixins/base.py (31 行)", "TraeSoloProtocol——所有 Mixin 接口的类型协议"],
        ["📡 CDP 通信层", "cdp_client.py (693 行)", "WebSocket 连接管理、eval_js() 核心方法、截图、DOM/Target/Network 操作"],
        ["📡 CDP 通信层", "connection_pool.py (220 行)", "连接池复用（10 上限）、空闲回收（300s）、存活检查"],
        ["🧩 辅助模块", "response_waiter.py (542 行)", "ResponseWaiter 状态机——9 步轮询策略、hash 稳定检测、自适应间隔"],
        ["🧩 辅助模块", "js_templates/ 6 模块 (927 行)", "按 Mixin 分组的 JS 模板函数（eval_js 注入字符串集中管理）"],
        ["🧩 辅助模块", "workspace_manager.py (388 行)", "项目类型检测 / Skills 推荐 / MCP Server 配置管理"],
        ["🧩 辅助模块", "project_manager.py (168 行)", "任务分析、角色推荐、多 Agent 计划执行"],
    ],
    col_widths=[2.5, 4.5, 9],
)

doc.add_heading("实际调用链路", level=2)

p = doc.add_paragraph()
run = p.add_run("submit 命令示例：【用户的 submit 命令】→ Typer 解析 → solo_session(连接池) → ChatMixin.type_message + TaskMixin.submit_task")
run.font.size = Pt(9)
run.font.name = "Consolas"
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

add_step_table([
    ["1", "cli.py → Typer 解析用户命令和参数"],
    ["2", "_run() → solo_session() → ConnectionPool.acquire() → CDPClient.connect()"],
    ["3", "TraeSoloController.Mixin 方法 → 调用 CDPClient.eval_js('document.querySelector().click()')"],
    ["4", "Runtime.evaluate → CDP WebSocket → Trae CN (Electron) 浏览器执行 JS"],
    ["5", "JS → DOM API → React 组件响应 → 操作结果通过 CDP 返回"],
    ["6", "StandardResponse → _display() → JSON / ndjson / --human 文本"],
])

doc.add_heading("工程统计", level=2)
add_table(
    ["指标", "数值"],
    [
        ["源码总行数", "6,116 行（src/traectl/，36 个 .py 文件）"],
        ["CLI 子模块", "7 个（_core / _query / _submit / _interact / _agent / _introspect / _workspace）"],
        ["Mixin 总行数", "1,178 行（8 个 Mixin + 1 个 Protocol 基类）"],
        ["JS 模板总行数", "927 行（6 个模块，从 Mixin 中统一抽取而来）"],
        ["CDPClient 行数", "693 行（单文件最大）"],
        ["测试总行数", "1,427 行（6 个测试文件）"],
        ["测试通过数", "172 个（pytest, 6.3s 全部通过）"],
        ["命令总数", "34 个"],
        ["DOM 选择器数量", "~30 个"],
        ["弹窗分类", "8 种"],
    ],
    col_widths=[3, 13],
)

# ══════════════════════════════════════════════════════════════════
# 3. 📦 模块职责一览
# ══════════════════════════════════════════════════════════════════
doc.add_heading("3. 📦 模块职责一览", level=1)
doc.add_paragraph("以下表格列出 traectl-CLI 全部核心模块及其一句话职责描述。")

add_table(
    ["模块", "行数", "一句话"],
    [
        ["cli/ (7 子模块)", "1,177", "Typer 命令定义、参数解析、输出格式化（JSON/NDJSON/--human）、高风险保护"],
        ["cdp_client.py", "693", "CDP WebSocket 通信底层：连接/重连/evalJS/截图/DOM/鼠标/网络操作"],
        ["js_templates/ (6 模块)", "927", "所有 JS DOM 查询模板函数（每个函数返回一段可注入的 JS 代码字符串）"],
        ["response_waiter.py", "542", "等待 SOLO 响应的状态机（9 步轮询策略：排队/哈希稳定/文件 mtime/超时处理）"],
        ["workspace_manager.py", "388", "工作区配置管理：项目类型检测、Skills 推荐、MCP Server 配置"],
        ["config.py", "235", "集中配置 + 常量 + DOM 选择器 + Agent 角色 + 环境变量覆写"],
        ["connection_pool.py", "220", "连接池复用、空闲 5 分钟回收、最大连接数限制、存活检查"],
        ["chat_mixin.py", "293", "聊天交互：新建/输入/发送/状态检测/hash 稳定/排队/自动确认"],
        ["task_mixin.py", "189", "任务管理：submit/ensure_idle/角色化提交/设置/模式切换/文件操作"],
        ["media_mixin.py", "187", "截图 + 重新生成 + 关闭弹窗 + 8 类弹窗自动处理"],
        ["project_manager.py", "168", "任务分析 + 角色推荐 + 多 Agent 计划执行"],
        ["response.py", "176", "标准化响应外壳：StandardResponse + JsonResponse + ok/error/dry_run 工厂"],
        ["git_mixin.py", "131", "Git 操作：status/stage/commit/diff/log/branch（通过本地 subprocess 执行）"],
        ["model_mixin.py", "126", "模型操作：列出可用模型、JS 注入切换当前模型"],
        ["terminal_mixin.py", "83", "终端操作：切换面板（Ctrl+`）、逐字符命令输入、读取 xterm 内容"],
        ["health_mixin.py", "71", "健康检查：CDP 连接状态 + SOLO 面板就绪 + DISPLAY 环境"],
        ["editor_mixin.py", "67", "编辑器状态读取：活跃文件、变更列表（diff）、接受/拒绝变更、半成品捞回"],
        ["controller.py", "36", "TraeSoloController——组合 8 个 Mixin 的入口组合类"],
        ["base.py", "31", "TraeSoloProtocol——所有 Mixin 接口的 typing.Protocol 基类"],
    ],
    col_widths=[4, 1.5, 10.5],
)

# ══════════════════════════════════════════════════════════════════
# 4. 💻 CLI 命令参考
# ══════════════════════════════════════════════════════════════════
doc.add_heading("4. 💻 CLI 命令参考", level=1)
doc.add_paragraph("以下按功能分组列出全部 34 个 CLI 命令。默认输出 JSON，加 --human 切换可读格式。")

doc.add_heading("📦 任务控制", level=2)
add_table(
    ["命令", "说明"],
    [
        ["submit", "向 SOLO 提交编码任务（支持 --role、--model、--no-wait）"],
        ["new", "创建新的空任务会话"],
        ["send", "发送消息到当前聊天（支持 --wait）"],
        ["status", "获取 SOLO 当前执行状态"],
        ["chat", "读取聊天记录（支持 --max-length）"],
        ["action get_tasks", "列出所有任务"],
        ["action switch_task -i N", "切换到指定任务"],
        ["action open_file <path>", "快速打开文件（Ctrl+P）"],
        ["action open_settings", "打开设置页面（命令面板）"],
        ["action open_mcp", "打开 MCP 配置"],
        ["action toggle_auto", "切换自动/手动模式"],
        ["action confirm", "自动点击确认弹窗"],
        ["action stop", "停止生成（⚠️ 高风险，需 --yes）"],
        ["action delete_task", "删除当前任务（⚠️ 高风险，需 --yes）"],
        ["regenerate", "重新生成最后一条回复（⚠️ 需 --yes）"],
        ["plan <json>", "多 Agent 计划执行"],
    ],
    col_widths=[5, 11],
)

doc.add_heading("🤖 模型管理", level=2)
add_table(
    ["命令", "说明"],
    [
        ["models", "列出所有可用模型及当前选中"],
        ["switch <model>", "切换模型（JS 注入打开选择器 → 遍历 → 点击）"],
        ["roles", "列出所有 Agent 角色及推荐模型"],
        ["analyze <task>", "分析任务推荐最佳角色和模型"],
    ],
    col_widths=[5, 11],
)

doc.add_heading("💬 聊天交互", level=2)
add_table(
    ["命令", "说明"],
    [
        ["chat", "读取聊天记录"],
        ["roles", "列出 Agent 角色"],
        ["regenerate", "重新生成最后回复"],
    ],
    col_widths=[5, 11],
)

doc.add_heading("📁 文件变化", level=2)
add_table(
    ["命令", "说明"],
    [
        ["editor", "获取编辑器状态及打开文件列表"],
        ["file-changes", "获取 SOLO 提议的文件变更（diff）"],
        ["accept", "接受文件变更（--file-path 指定单个，留空全部）"],
        ["reject", "拒绝文件变更（--file-path 指定单个，留空全部）"],
        ["file-status", "监视目录文件变化状态（--dir, --glob, --fields）"],
    ],
    col_widths=[5, 11],
)

doc.add_heading("💻 终端操作", level=2)
add_table(
    ["命令", "说明"],
    [
        ["terminal", "切换终端面板显示/隐藏（Ctrl+`）"],
        ["exec <command>", "在 xterm 中逐字符输入并执行命令"],
        ["terminal-content", "获取终端面板文本内容"],
    ],
    col_widths=[5, 11],
)

doc.add_heading("📸 媒体与弹窗", level=2)
add_table(
    ["命令", "说明"],
    [
        ["screenshot", "截取界面截图（--save-path）"],
        ["close-dialog", "关闭当前弹窗"],
        ["auto-recover", "智能扫描并处理页面上所有弹窗"],
    ],
    col_widths=[5, 11],
)

doc.add_heading("🔧 Git", level=2)
add_table(
    ["命令", "说明"],
    [
        ["git status", "查看仓库状态"],
        ["git stage --file-path <path>", "暂存文件（留空全部暂存）"],
        ["git commit -m <msg>", "提交变更"],
        ["git diff", "查看变更差异"],
        ["git log", "查看提交历史"],
        ["git branch", "查看分支"],
    ],
    col_widths=[5, 11],
)

doc.add_heading("🩺 诊断与自省", level=2)
add_table(
    ["命令", "说明"],
    [
        ["health", "健康检查（CDP 连接 + SOLO 面板 + DISPLAY 环境）"],
        ["commands", "列出所有可用命令"],
        ["schema [--command <name>]", "输出命令 JSON Schema"],
        ["categories", "按类别分组列出所有命令"],
        ["exit-codes", "列出退出码说明"],
        ["help [command]", "显示命令帮助"],
        ["version", "输出版本信息"],
    ],
    col_widths=[5, 11],
)

doc.add_heading("🛠️ 维护操作", level=2)
add_table(
    ["命令", "说明"],
    [
        ["action stop", "停止生成（⚠️）"],
        ["action delete_task", "删除当前任务（⚠️）"],
        ["action open_settings", "打开设置"],
        ["action open_mcp", "打开 MCP 配置"],
        ["action toggle_auto", "切换自动模式"],
        ["action confirm", "确认弹窗"],
        ["action open_file <path>", "打开文件"],
        ["install-skills", "安装 traectl Agent Skills"],
    ],
    col_widths=[5, 11],
)

doc.add_heading("⚙️ 配置管理", level=2)
add_table(
    ["命令", "说明"],
    [
        ["config get <key>", "读取配置项"],
        ["config set <key> <value>", "设置配置项"],
        ["config list", "列出所有配置"],
        ["config export", "导出配置为 JSON"],
    ],
    col_widths=[5, 11],
)

doc.add_heading("📂 工作区管理", level=2)
add_table(
    ["命令", "说明"],
    [
        ["workspace init --path <dir>", "初始化项目工作区（自动检测项目类型）"],
        ["workspace setup-mcp <name>", "配置 MCP 服务器（--command, --args）"],
        ["install-skills", "安装 Agent Skills（--target, --global）"],
    ],
    col_widths=[5, 11],
)

# ══════════════════════════════════════════════════════════════════
# 5. 🔄 核心流程
# ══════════════════════════════════════════════════════════════════
doc.add_heading("5. 🔄 核心流程", level=1)

doc.add_heading("5.1 任务提交流程（submit）", level=2)
doc.add_paragraph("以下步骤展示从输入 prompt 到获取 SOLO 响应的完整链路：")
add_step_table([
    ["1", "用户执行 traectl submit \"任务描述\" --role backend --no-wait"],
    ["2", "Typer 解析参数，进入 submit() 命令函数"],
    ["3", "submit() 检查 --no-wait；若否，先调用 ensure_idle() 等待上一个任务完成"],
    ["4", "_run() 创建 solo_session → ConnectionPool.acquire() → CDPClient.connect()"],
    ["5", "若指定 --role，调用 submit_task_with_role()：查 AGENT_ROLES → 推荐模型 → switch_model() → 组装 prompt"],
    ["6", "start_new_task()：CDP eval_js('document.querySelector(new-task-btn).click()')"],
    ["7", "type_message()：JS focus_input_box → CDP Input.insertText 输入文本"],
    ["8", "send_message()：CDP eval_js('document.querySelector(send-btn).click()')"],
    ["9", "ResponseWaiter.wait(timeout, retries)：状态机 9 步轮询等待（排队→生成→hash稳定→mtime→完成）"],
    ["10", "连接归还 ConnectionPool，结果通过 mk_ok() 包装为标准 JSON 输出"],
])

doc.add_heading("5.2 自动弹窗处理流程（auto-recover）", level=2)
doc.add_paragraph("auto-recover 命令自动扫描并处理页面上出现的各类弹窗，支持 8 种弹窗类型：")
add_step_table([
    ["1", "CDP eval_js(scan_dialogs()) 扫描页面所有可见弹窗，提取文本和按钮"],
    ["2", "用 DIALOG_PATTERNS 正则匹配树分类（匹配 pattern 最多的类型胜出）"],
    ["3", "server_error / error_2000000 → 点击「继续」，最多重试 3 次"],
    ["4", "update_dialog → 点击「忽略」/「稍后」"],
    ["5", "queue_reminder → 读取排队位置号，仅报告"],
    ["6", "confirm_dialog → 自动点击「运行」/「确认」"],
    ["7", "permission_dialog / resource_warning / retry_dialog → 针对性处理"],
    ["8", "未知弹窗 → 退回到 close_dialog：关闭按钮或 Escape 键"],
])

doc.add_heading("5.3 模型切换流程（关键 JS 注入场景）", level=2)
add_step_table([
    ["1", "get_current_model()：JS 读取选择器显示文本，已是目标模型则跳过"],
    ["2", "trigger_model_selector()：dispatchEvent(new MouseEvent('mousedown', {bubbles: true}))"],
    ["3", "轮询 17 次 × 0.3s 等待列表渲染（count_model_items 检测）"],
    ["4", "trigger_model_selector_js()：鼠标事件失败则用纯 JS 方式重试"],
    ["5", "click_model_item()：querySelectorAll 遍历匹配 → scrollIntoView → .click()"],
    ["6", "验证：get_current_model() → _model_name_match() 模糊匹配（处理 Beta 后缀）"],
    ["7", "兜底：SQLite 覆写 state.vscdb + 重启 Trae CN"],
])

doc.add_heading("5.4 响应等待状态机（ResponseWaiter）", level=2)
doc.add_paragraph("ResponseWaiter 是 submit/send 命令的核心等待引擎，542 行独立模块。其状态转换如下：")
add_table(
    ["状态", "触发条件", "动作"],
    [
        ["初始化", "开始等待", "记录起始时间，清空心跳计数"],
        ["自动确认", "检测到 inline_delete 弹窗", "自动点击确认按钮"],
        ["生成检测", "is_generating = True", "标记 saw_generating，开始稳定计数"],
        ["生成完成", "saw_generating True→False", "最强完成信号，验证最后角色为 assistant"],
        ["排队检测", "检测到排队面板", "读取排队位置，每 30s 更新"],
        ["哈希稳定", "连续 STABLE_THRESHOLD（3）次哈希一致", "进入完成准备阶段"],
        ["文件稳定", "连续 STABLE_THRESHOLD 次文件 mtime 不变", "返回最终内容"],
        ["60% 提前返回", "stable_count≥1 + saw_generating + elapsed>60%", "提前返回已有内容（优化体验）"],
        ["超时处理", "超过 timeout", "提取聊天内容 + 编辑器半成品 + 面板快照 → 重试或返回超时"],
    ],
    col_widths=[2.5, 5, 8.5],
)

# ══════════════════════════════════════════════════════════════════
# 6. 📡 CDP 通信机制
# ══════════════════════════════════════════════════════════════════
doc.add_heading("6. 📡 CDP 通信机制", level=1)

doc.add_heading("6.1 连接架构", level=2)
doc.add_paragraph("CDP（Chrome DevTools Protocol）是 traectl 与 Trae CN 浏览器实例通信的核心协议。通过 WebSocket 连接 Chrome DevTools，发送 Runtime.evaluate 等命令操控页面。")

doc.add_paragraph("端口与连接配置：")
add_table(
    ["配置项", "环境变量", "默认值", "说明"],
    [
        ["CDP 主机", "TRAECTL_CDP_HOST", "127.0.0.1", "Trae CN 浏览器的 CDP 监听地址"],
        ["CDP 端口", "TRAECTL_CDP_PORT", "9222", "Trae CN 浏览器的 CDP 调试端口"],
        ["SOLO 超时", "TRAECTL_TIMEOUT", "300", "等待 SOLO 响应的超时秒数"],
        ["稳定阈值", "TRAECTL_STABLE_THRESHOLD", "3", "hash 或文件 mtime 稳定连续次数"],
        ["轮询间隔", "TRAECTL_POLL_INTERVAL", "2", "响应轮询间隔秒数"],
        ["最大重试次数", "TRAECTL_CDP_MAX_RETRIES", "5", "CDP 连接最大重试次数"],
        ["初始重试间隔", "TRAECTL_CDP_INITIAL_RETRY_INTERVAL", "1", "指数退避初始间隔"],
        ["最大重试间隔", "TRAECTL_CDP_MAX_RETRY_INTERVAL", "30", "指数退避最大间隔"],
    ],
    col_widths=[3, 5, 2, 6],
)

doc.add_heading("6.2 页面匹配策略", level=2)
doc.add_paragraph("从 CDP 的页面列表中匹配目标页面时采用三级降级策略：")
add_table(
    ["优先级", "匹配条件", "示例"],
    [
        ["1（最高）", "页面 URL 包含 vscode-file 关键字", "vscode-file://vscode-app/..."],
        ["2", "页面标题包含 \"Trae\" 字样", "Trae CN - ..."],
        ["3（兜底）", "任意可用的 page 类型", "第一个非空的 page target"],
    ],
    col_widths=[2, 5, 9],
)

doc.add_heading("6.3 心跳与重连", level=2)
doc.add_paragraph("为保持连接稳定，CDP 通信层实现了心跳和自动重连机制：")
add_table(
    ["机制", "参数", "行为"],
    [
        ["Keepalive Ping", "每 3 秒", "通过 websocket ping 检查连接存活"],
        ["指数退避重连", "1s → 2s → 4s → ... → 30s（上限）", "断连后按指数增长间隔尝试重连"],
        ["最大重试次数", "5 次", "超过后抛出连接异常"],
        ["连接安全检测", "is_alive()", "每次从连接池取连接时验证存活状态"],
    ],
    col_widths=[3, 4, 9],
)

doc.add_heading("6.4 主要 CDP 命令", level=2)
doc.add_paragraph("traectl 使用以下 CDP 命令与 Trae CN 浏览器交互：")
add_table(
    ["CDP 方法", "目的", "使用场景"],
    [
        ["Runtime.evaluate", "在目标页面中执行任意 JavaScript", "所有 DOM 操作、状态查询、点击事件（核心方法）"],
        ["Page.captureScreenshot", "截取页面当前视图", "screenshot 命令"],
        ["Input.insertText", "向焦点元素输入文本", "type_message——React contenteditable 的唯一可靠方式"],
        ["Input.dispatchKeyEvent", "模拟键盘按键", "Ctrl+` 切换终端、逐字符命令输入、快捷键操作"],
        ["Input.dispatchMouseEvent", "模拟鼠标事件", "兜底方案（React 合成事件可能拦截）"],
        ["Page.navigate", "导航到指定 URL", "打开/刷新页面"],
        ["DOM.querySelector", "在 DOM 中查找元素", "定位聊天框、按钮、编辑器等元素"],
        ["Target.activateTarget", "激活指定页面标签", "在多个页面间切换焦点"],
    ],
    col_widths=[3.5, 5, 7.5],
)

doc.add_heading("6.5 连接池机制", level=2)
add_table(
    ["特性", "说明"],
    [
        ["连接复用", "按 (host, port) 键复用已有空闲连接，命令执行完归还而非断开"],
        ["最大连接数", "默认 10，超限时回收最久未使用的空闲连接"],
        ["空闲超时", "默认 300s（5 分钟），后台每 60s 扫描回收超时空闲连接"],
        ["存活检查", "acquire 时 ping 检测，死亡连接自动移除并新建"],
        ["线程安全", "asyncio.Lock 保护所有池操作"],
        ["全局单例", "get_pool() 获取全局连接池实例"],
    ],
    col_widths=[3, 13],
)

# ── 保存 ──────────────────────────────────────────────────────────
output_path = "/home/ubuntu/traectl-CLI/docs/traectl-架构文档.docx"
doc.save(output_path)
import os
size = os.path.getsize(output_path)
print(f"✅ 文档已保存: {output_path} ({size} bytes)")
