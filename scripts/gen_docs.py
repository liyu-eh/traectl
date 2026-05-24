#!/usr/bin/env python3
"""生成 traectl-CLI 技术栈 + 架构文档 (.docx)。"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = "/home/ubuntu/traectl-CLI/docs"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "traectl-架构文档.docx")

doc = Document()

# ── 样式 ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x5C, 0xB0)

def set_cell_shading(cell, color="D9E2F3"):
    """设置单元格底色 (Light Shading Accent 1 风格)。"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table(headers, rows, col_widths=None):
    """添加带样式的表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # 表后间距
    return table


# ═══════════════════════════════════════════════════════════════════
# 封面 / 标题
# ═══════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("traectl-CLI 技术栈与架构文档")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1A, 0x5C, 0xB0)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("版本 3.1.0 · Python ≥3.11 · 基于 Chrome DevTools Protocol")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
# 1. 📋 技术栈
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("1. 📋 技术栈", level=1)
doc.add_paragraph("以下表格列出 traectl-CLI 实际采用的核心技术及其用途。")

add_table(
    ["技术", "版本", "用途"],
    [
        ["Python", "≥3.11", "项目运行语言，充分利用 asyncio 异步 IO"],
        ["Typer (Click)", "≥0.12", "CLI 框架：参数解析、子命令注册、自动帮助文档"],
        ["websockets", "≥12", "CDP WebSocket 全双工通信，与 Chrome DevTools 协议交互"],
        ["Rich", "≥13", "人类可读输出：stderr 日志、Syntax 代码高亮、面板渲染"],
        ["pytest", "≥7", "单元测试框架，覆盖 CLI、CDP、配置、工作区等模块"],
        ["asyncio", "stdlib", "异步 IO 驱动：CDP 通信、连接池、响应轮询"],
        ["setuptools", "≥64", "项目打包与 editable install（pyproject.toml 声明）"],
    ],
    col_widths=[3.5, 2.5, 9],
)

doc.add_paragraph("注：pyproject.toml 中显式依赖只有 typer 和 rich；websockets 为运行时直接导入，pytest 为测试依赖。")

# ═══════════════════════════════════════════════════════════════════
# 2. 🏗️ 三层架构
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("2. 🏗️ 三层架构", level=1)
doc.add_paragraph(
    "traectl-CLI 采用清晰的三层架构设计，从顶层 CLI 到底层 CDP 通信逐层隔离职责："
)

add_table(
    ["层", "模块", "职责"],
    [
        [
            "🖥️ CLI 层",
            "cli.py",
            "命令解析（Typer）、输出格式化（JSON/NDJSON/--human）、退出码管理、公共标志注入",
        ],
        [
            "🎮 Controller 层",
            "controller.py + 7 个 Mixin",
            "业务逻辑：发送消息、切换模型、管理任务、操作编辑器/终端/Git/截图/弹窗",
        ],
        [
            "📡 CDP 通信层",
            "cdp_client.py + connection_pool.py",
            "WebSocket 连接管理、CDP 命令发送/响应匹配、连接池复用与空闲回收",
        ],
    ],
    col_widths=[3, 5, 7],
)

doc.add_heading("实际调用链路", level=2)
p = doc.add_paragraph()
p.add_run("submit 命令示例：").bold = True
code = p.add_run(
    "\n  cli.py submit → solo_session(连接池) → TaskMixin.submit_task\n"
    "               → ChatMixin.type_message + send_message\n"
    "               → ResponseWaiter.wait → CDPClient.send_command(Runtime.evaluate)\n"
    "               → WebSocket → Trae CN Chrome"
)
code.font.name = 'Consolas'
code.font.size = Pt(9)

# ═══════════════════════════════════════════════════════════════════
# 3. 📦 模块职责
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("3. 📦 模块职责一览", level=1)
doc.add_paragraph("以下表格列出 traectl-CLI 全部核心模块及其一句话职责描述。")

add_table(
    ["模块", "行数", "一句话"],
    [
        ["cli.py", "1171", "Typer 命令解析，全部 CLI 入口（所有子命令定义 + 公共标志 + 输出格式化）"],
        ["cdp_client.py", "693", "CDP WebSocket 通信底层：连接/重连/evalJS/截图/DOM 操作/鼠标/网络"],
        ["js_templates.py", "927", "所有 JS DOM 查询模板函数（每个函数返回一段可注入的 JS 字符串）"],
        ["response_waiter.py", "542", "等待 SOLO 响应的状态机（轮询生成/排队/哈希/文件 mtime 状态）"],
        ["config.py", "235", "集中配置 + 常量 + DOM 选择器 + Agent 角色 + 配置持久化"],
        ["chat_mixin.py", "293", "聊天交互：输入/发送/状态检测/排队/自动确认/任务列表查询"],
        ["task_mixin.py", "189", "任务管理：submit/ensure_idle/模式切换/打开文件"],
        ["media_mixin.py", "187", "截图 + 重新生成 + 关闭弹窗 + 自动弹窗处理"],
        ["workspace_manager.py", "388", "项目类型检测 + Skills 推荐 + MCP Server 配置管理"],
        ["connection_pool.py", "220", "连接池：按 (host,port) 复用、空闲 5 分钟回收、最大连接数限制"],
        ["project_manager.py", "168", "任务分析 + 角色推荐 + 多 Agent 计划执行"],
        ["model_mixin.py", "126", "模型操作：列出可用模型、切换当前模型"],
        ["editor_mixin.py", "67", "编辑器操作：获取活跃文件信息、文件变更列表、接受/拒绝变更"],
        ["git_mixin.py", "131", "Git 操作：status/stage/commit/diff/log/branch（通过终端命令执行）"],
        ["terminal_mixin.py", "83", "终端操作：切换面板、执行命令、获取终端内容"],
        ["response.py", "147", "标准化响应外壳：StandardResponse + JsonResponse + ok/error/dry_run_plan"],
        ["health_mixin.py", "71", "健康检查：获取 CDP 连接状态、页面存活、SOLO 响应性"],
        ["base.py", "31", "TraeSoloProtocol — typing.Protocol 接口定义（静态类型检查）"],
        ["controller.py", "36", "组合 7 个 Mixin 的入口类 TraeSoloController"],
    ],
    col_widths=[3.5, 1.5, 10],
)

# ═══════════════════════════════════════════════════════════════════
# 4. 💻 CLI 命令参考
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("4. 💻 CLI 命令参考", level=1)
doc.add_paragraph("以下按功能分组列出所有 CLI 命令。")

command_groups = [
    ("📦 任务控制", [
        ("submit", "向 SOLO 提交编码任务", "traectl submit \"实现登录功能\""),
        ("new", "创建新的空任务会话", "traectl new"),
        ("action get_tasks", "获取当前任务列表", "traectl action get_tasks"),
        ("action switch_task", "切换到指定任务", "traectl action switch_task --task-index 2"),
        ("action delete_task", "删除当前任务", "traectl action delete_task"),
        ("action stop", "停止生成", "traectl action stop"),
        ("send", "发送消息到当前聊天", "traectl send \"修改样式\" --wait"),
        ("plan", "多 Agent 计划执行", "traectl plan '[{\"role\":\"architect\",\"task\":\"设计API\"}]'"),
    ]),
    ("🤖 模型管理", [
        ("models", "列出可用模型", "traectl models"),
        ("switch", "切换模型", "traectl switch DeepSeek-V4-Pro"),
    ]),
    ("💬 聊天交互", [
        ("chat", "读取聊天记录", "traectl chat --max-length 3000"),
        ("roles", "列出 Agent 角色", "traectl roles"),
        ("regenerate", "重新生成最后回复", "traectl regenerate"),
    ]),
    ("📁 文件变化", [
        ("editor", "获取编辑器状态", "traectl editor"),
        ("file-changes", "列出文件变更", "traectl file-changes"),
        ("accept", "接受文件变更", "traectl accept --file-path src/main.py"),
        ("reject", "拒绝文件变更", "traectl reject"),
        ("file-status", "文件状态监视", "traectl file-status --dir . --glob \"*.py\""),
    ]),
    ("💻 终端操作", [
        ("terminal", "切换终端面板", "traectl terminal"),
        ("exec", "执行终端命令", "traectl exec \"npm run build\""),
        ("terminal-content", "获取终端内容", "traectl terminal-content"),
    ]),
    ("📸 媒体与弹窗", [
        ("screenshot", "截取界面截图", "traectl screenshot --save-path ./shot.png"),
        ("close-dialog", "关闭弹窗", "traectl close-dialog"),
        ("auto-recover", "自动弹窗处理", "traectl auto-recover"),
    ]),
    ("🔧 Git", [
        ("git status", "查看 Git 状态", "traectl git status"),
        ("git stage", "暂存文件", "traectl git stage --file-path src/main.py"),
        ("git commit", "提交变更", "traectl git commit -m \"fix: login bug\""),
        ("git diff", "查看变更 diff", "traectl git diff"),
        ("git log", "查看提交历史", "traectl git log"),
        ("git branch", "查看分支", "traectl git branch"),
    ]),
    ("🩺 诊断与自省", [
        ("status", "获取 SOLO 状态", "traectl status"),
        ("health", "健康检查", "traectl health"),
        ("version", "输出版本信息", "traectl version"),
        ("commands", "列出所有命令", "traectl commands"),
        ("help", "命令帮助", "traectl help submit"),
        ("schema", "命令 JSON Schema", "traectl schema --command submit"),
        ("exit-codes", "退出码说明", "traectl exit-codes"),
    ]),
    ("🛠️ 维护操作", [
        ("action open_settings", "打开设置面板", "traectl action open_settings"),
        ("action open_mcp", "打开 MCP 设置", "traectl action open_mcp"),
        ("action toggle_auto", "切换自动模式", "traectl action toggle_auto --disable-auto"),
        ("action confirm", "确认对话框", "traectl action confirm"),
        ("action open_file", "打开文件", "traectl action open_file --file-path main.py"),
        ("install-skills", "安装 Agent Skill", "traectl install-skills --target ./skills"),
    ]),
    ("⚙️ 配置管理", [
        ("config get", "读取配置", "traectl config get cdp_port"),
        ("config set", "设置配置", "traectl config set solo_timeout 600"),
        ("config list", "列出所有配置", "traectl config list"),
        ("config export", "导出配置 JSON", "traectl config export"),
    ]),
    ("📂 工作区管理", [
        ("workspace init", "初始化工作区", "traectl workspace init --path ./my-project"),
        ("workspace setup-mcp", "配置 MCP", "traectl workspace setup-mcp --provider deepseek"),
    ]),
    ("📋 任务分析", [
        ("analyze", "分析任务推荐角色", "traectl analyze \"实现一个 REST API\""),
    ]),
]

for group_name, commands in command_groups:
    doc.add_heading(group_name, level=2)
    add_table(
        ["命令", "说明", "示例"],
        [[c[0], c[1], c[2]] for c in commands],
        col_widths=[4, 5, 6],
    )

# ═══════════════════════════════════════════════════════════════════
# 5. 🔄 核心流程
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("5. 🔄 核心流程", level=1)

doc.add_heading("5.1 任务提交流程（submit）", level=2)
doc.add_paragraph("以下步骤展示从输入 prompt 到获取 SOLO 响应的完整链路：")

add_table(
    ["步骤", "操作", "说明"],
    [
        ["1", "输入 prompt", "用户输入任务描述到 CLI（如 \"实现登录功能\"）"],
        ["2", "ensure_idle", "检测当前任务是否空闲；若繁忙则等待完成或报错"],
        ["3", "start_new_task", "点击「新建任务」按钮创建新会话"],
        ["4", "type_message", "使用 CDP Input.insertText 将 prompt 注入富文本输入框"],
        ["5", "send_message", "点击发送按钮提交消息"],
        ["6", "ResponseWaiter.wait", "状态机轮询等待：检测排队 → 生成 → 哈希稳定 → 文件 mtime → 完成"],
        ["7", "返回响应", "JSON 格式返回 SOLO 生成的代码/文本/文件变更信息"],
    ],
    col_widths=[1, 3.5, 10.5],
)

doc.add_heading("5.2 自动弹窗处理流程（auto-recover）", level=2)
doc.add_paragraph("auto-recover 命令自动扫描并处理页面上出现的各类弹窗：")

add_table(
    ["步骤", "操作", "说明"],
    [
        ["1", "扫描弹窗", "通过 DOM 查询检测 overlay/modal/mask/backdrop 等弹窗容器"],
        ["2", "提取文字", "读取弹窗内的文本内容"],
        ["3", "匹配模式", "与 DIALOG_PATTERNS 中的关键词进行匹配"],
        ["4", "分发处理", "按类型执行不同操作：服务端异常 → 继续重试；排队 → 等待；确认 → 点击确认"],
        ["5", "记录日志", "记录处理结果到 progress log"],
    ],
    col_widths=[1, 3.5, 10.5],
)

doc.add_heading("5.3 响应等待状态机（ResponseWaiter）", level=2)
doc.add_paragraph("ResponseWaiter 是 submit/send 命令的核心等待引擎，其状态转换如下：")

add_table(
    ["状态", "触发条件", "动作"],
    [
        ["初始化", "开始等待", "记录起始时间，清空心跳计数"],
        ["排队检测", "队列非空", "记录排队信息，继续轮询"],
        ["生成检测", "is_generating = true", "标记 saw_generating，开始稳定计数"],
        ["哈希稳定", "连续 STABLE_THRESHOLD 次哈希一致", "进入文件变更检测阶段"],
        ["文件稳定", "连续 STABLE_THRESHOLD 次文件 mtime 不变", "返回最终内容"],
        ["超时处理", "超过 timeout", "触发重试（最多 max_retries 次），或返回超时错误"],
    ],
    col_widths=[2.5, 4.5, 8],
)

# ═══════════════════════════════════════════════════════════════════
# 6. 📡 CDP 通信机制
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("6. 📡 CDP 通信机制", level=1)

doc.add_heading("6.1 连接架构", level=2)
doc.add_paragraph("CDP（Chrome DevTools Protocol）是 traectl 与 Trae CN 浏览器实例通信的核心协议。")

add_table(
    ["配置项", "默认值", "说明"],
    [
        ["CDP 端口", "9222 (Hermes) / 9224 (Trae CN)", "两个独立的 Chrome 实例各有自己的调试端口"],
        ["连接池复用", "按 (host, port) 组合", "同一 (host, port) 对复用同一个 WebSocket 连接"],
        ["空闲超时", "300 秒（5 分钟）", "超过此时间未使用的连接将被自动回收关闭"],
        ["最大连接数", "10", "连接池同时保持的最大连接数量"],
    ],
    col_widths=[3.5, 5, 6.5],
)

doc.add_heading("6.2 页面匹配策略", level=2)
doc.add_paragraph("从 CDP 的页面列表中匹配目标页面时采用三级降级策略：")

add_table(
    ["优先级", "匹配条件", "示例"],
    [
        ["1（最高）", "页面 URL 包含 vscode-file 关键字", "vscode-file://vscode-app/..."],
        ["2", "页面标题包含 \"Trae\" 字样", "Trae CN - ..."],
        ["3（兜底）", "任意可用的页面", "第一个非空的 page target"],
    ],
    col_widths=[2.5, 5, 7.5],
)

doc.add_heading("6.3 心跳与重连", level=2)
doc.add_paragraph("为保持连接稳定，CDP 通信层实现了心跳和自动重连机制：")

add_table(
    ["机制", "参数", "行为"],
    [
        ["Keepalive Ping", "每 3 秒", "通过 websocket ping 检查连接存活"],
        ["指数退避重连", "1s → 2s → 4s → ... → 30s (max)", "断连后按指数增长间隔尝试重连，上限 30 秒"],
        ["最大重试次数", "5 次", "超过后抛出连接异常"],
        ["连接安全检测", "is_alive()", "每次从连接池取连接时验证存活状态"],
    ],
    col_widths=[3.5, 3.5, 8],
)

doc.add_heading("6.4 主要 CDP 命令", level=2)
doc.add_paragraph("traectl 使用以下 CDP 命令与 Trae CN 浏览器交互：")

add_table(
    ["CDP 方法", "目的", "使用场景"],
    [
        ["Runtime.evaluate", "在目标页面中执行任意 JavaScript", "所有 DOM 操作、状态查询、点击事件"],
        ["Page.captureScreenshot", "截取页面当前视图", "screenshot 命令"],
        ["Input.insertText", "向焦点元素输入文本", "type_message — React contenteditable 的唯一可靠方式"],
        ["Page.navigate", "导航到指定 URL", "打开/刷新页面"],
        ["DOM.querySelector", "在 DOM 中查找元素", "定位聊天框、按钮等元素"],
    ],
    col_widths=[4, 4, 7],
)

# ── 最后保存 ─────────────────────────────────────────────────────
os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_FILE)
print(f"✅ 文档已生成: {OUTPUT_FILE}")
print(f"   大小: {os.path.getsize(OUTPUT_FILE)} bytes")
